from pathlib import Path

from file_compare_app.analyzers.base import Analyzer, AnalyzerDependencyError
from file_compare_app.core.models import Change, ChangeLocation, ComparisonResult
from file_compare_app.core.text_diff import line_changes


class DocxAnalyzer(Analyzer):
    def compare(self, left_path: Path, right_path: Path) -> ComparisonResult:
        left_lines = _extract_docx_lines(left_path)
        right_lines = _extract_docx_lines(right_path)
        changes: list[Change] = []
        for index, (kind, left_index, right_index, before, after) in enumerate(line_changes(left_lines, right_lines), start=1):
            changes.append(
                Change(
                    change_id=f"docx-{index}",
                    change_type=kind,  # type: ignore[arg-type]
                    severity="normal",
                    source_location=ChangeLocation(str(left_path), line=left_index + 1) if left_index is not None else None,
                    target_location=ChangeLocation(str(right_path), line=right_index + 1) if right_index is not None else None,
                    before=before,
                    after=after,
                    format="docx",
                    confidence=1.0,
                    notes="paragraph/table text",
                )
            )
        return ComparisonResult(str(left_path), str(right_path), "docx", changes)


def _extract_docx_lines(path: Path) -> list[str]:
    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise AnalyzerDependencyError("Для сравнения DOCX установите python-docx.") from exc

    document = Document(path)
    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return lines
